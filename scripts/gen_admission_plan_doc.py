# -*- coding: utf-8 -*-
"""Sinh file Word mo ta ke hoach tinh nang auto-generate trang tinh diem tu PDF de an tuyen sinh."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x55, 0x55, 0x55)
CODE_BG = "F2F2F2"


def set_cell_bg(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
    # nen xam nhe
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)
    return p


def add_why(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    label = p.add_run("Vi sao: ")
    label.bold = True
    label.font.color.rgb = ACCENT
    label.font.size = Pt(10.5)
    body = p.add_run(text)
    body.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


doc = Document()

# Style mac dinh
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# ----- Trang bia -----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("KE HOACH PHAT TRIEN TINH NANG\nTu dong tao trang tinh diem theo tung truong tu PDF de an tuyen sinh")
trun.bold = True
trun.font.size = Pt(20)
trun.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = sub.add_run("Du an ZPATH - He thong huong nghiep AI")
srun.italic = True
srun.font.size = Pt(12)
srun.font.color.rgb = MUTED

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Tai lieu phan tich va de xuat phuong an ky thuat\nNgay lap: 09/06/2026").font.size = Pt(10.5)

doc.add_paragraph()

# ----- 0. Tom tat nhanh -----
h1(doc, "0. Tom tat nhanh (TL;DR)")
doc.add_paragraph(
    "Muc tieu: khi nhom them mot file PDF de an tuyen sinh cua mot truong, he thong "
    "tu sinh ra mot trang tinh diem rieng cho truong do; admin phai phe duyet truoc khi "
    "trang xuat hien tren web chinh."
)
add_bullets(doc, [
    ("Phuong an toi uu: ", "Hybrid = Config-driven + AI trich xuat ban nhap + admin review + publish bang co trong DB."),
    ("Muc do tu dong: ", "AI parse PDF ra CONFIG (JSON) nhap -> admin sua trong trang quan tri -> duyet. KHONG full-auto, KHONG sinh code TypeScript chay production."),
    ("Cach phe duyet/deploy: ", "Bat/tat bang co trang thai trong DB (draft -> pending_review -> published). Trang hien ngay sau khi duyet, KHONG can build/redeploy lai."),
    ("Tinh than: ", "Giu nguyen nguyen tac tach lop cua file demo-manual.pdf, nhung doi 'moi truong = mot module code' thanh 'moi truong = mot ban config du lieu'."),
])

doc.add_page_break()

# ----- 1. Boi canh & muc tieu -----
h1(doc, "1. Boi canh va muc tieu")
doc.add_paragraph(
    "Hien tai ZPATH la ung dung Next.js 16 + Supabase. He thong tinh diem xet tuyen "
    "(admission) dang duoc viet rieng cho tung truong bang code (HUST, FTU, UET, NEU, VINUNI). "
    "Moi truong la mot 'module' TypeScript viet tay, va giao dien tinh diem dung mot component "
    "trung tam co cac nhanh if (isFtu) / if (UET) / if (HUST)."
)
doc.add_paragraph("Yeu cau moi cua ban gom 3 phan:")
add_bullets(doc, [
    "Them file PDF de an tuyen sinh cua mot truong.",
    "He thong tu dong tao ra mot trang tinh diem rieng cho truong do.",
    "Nhom phai phe duyet truoc khi trang duoc deploy len web chinh.",
])

# ----- 2. Danh gia PDF -----
h1(doc, "2. Danh gia file demo-manual.pdf")
doc.add_paragraph(
    'File co ten "So do lam viec va quy uoc mo rong Admission Engines" (7 trang). '
    "Noi dung thuc chat la QUY UOC TO CHUC CODE de scale engine xet tuyen len 20+ truong "
    "ma it conflict khi nhieu nguoi cung lam."
)
h2(doc, "2.1. Tom tat noi dung PDF")
add_bullets(doc, [
    ("Tach 4 lop: ", "shared/ (that mong) -> schools/<school>/ (rule, benchmark, config rieng) -> engines/<engine>/ (schema + validate + scoring + mapping) -> UI container chi dieu phoi."),
    ("Contract chung: ", "schoolCode, engineId, year, payload, benchmark."),
    ("Registry mong: ", "chi map schoolCode -> module / engineId -> engine, khong chua logic."),
    ("Checklist: ", "khi them truong moi / engine moi, va danh sach file de conflict (registry trung tam, types chung, calculator mot-file khong lo)."),
])
h2(doc, "2.2. Nhan xet")
p = doc.add_paragraph()
p.add_run("Diem tot: ").bold = True
p.add_run("PDF dung ve nguyen tac ky thuat - tach lop ro rang, day rule rieng ra khoi UI, giu shared mong la cach lam ben vung.")
p = doc.add_paragraph()
p.add_run("Han che quan trong: ").bold = True
p.add_run(
    "PDF giai mot bai toan KHAC voi yeu cau cua ban. PDF gia dinh dev VIET TAY mot module code "
    "moi truong. Con ban muon UPLOAD PDF la co trang ngay. Hai viec nay khac han nhau ve trien khai."
)
p = doc.add_paragraph()
p.add_run("Mau thuan voi code hien tai: ").bold = True
p.add_run(
    "File src/components/admission/AdmissionCalculatorSection.tsx dang lam dung dieu PDF canh bao "
    "(component trung tam chua rule rieng, de conflict)."
)

# ----- 3. Van de cot loi -----
h1(doc, "3. Van de cot loi can quyet: Code-module hay Config-driven?")
doc.add_paragraph(
    "Day la quyet dinh anh huong toan bo kien truc. Hai trang thai khac nhau:"
)
add_bullets(doc, [
    ("Code-module (cach PDF de xuat): ", "moi truong la mot file code TypeScript. Manh, linh hoat, nhung KHONG the auto-sinh tu PDF (khong the vua upload PDF vua tu viet code an toan chay production)."),
    ("Config-driven (cach can cho tinh nang nay): ", "moi truong la mot ban CONFIG (JSON) mo ta cong thuc; mot engine tong doc config va tinh. Cho phep 'upload PDF -> sinh config -> chay duoc' ma khong dong vao code."),
])
add_why(doc,
    "Du lieu de an tuyen sinh rat rui ro ve do chinh xac (THPT, TSA, HSA, SAT, qui doi IELTS, "
    "diem uu tien khu vuc, cong thuc nhan he so khac nhau tung truong). Sinh code tu dong roi chay "
    "production la cuc ky nguy hiem. Nguoc lai, config la du lieu - de kiem tra, de sua, de rollback, "
    "khong can deploy. Vi vay tinh nang nay BUOC phai di huong config-driven."
)

# ----- 4. Phuong an toi uu -----
h1(doc, "4. Phuong an toi uu duoc khuyen nghi")
doc.add_paragraph(
    "Hybrid: Config-driven + AI trich xuat ban nhap + admin review + publish bang co DB."
)
add_bullets(doc, [
    ("Muc do tu dong: ", "AI parse PDF -> sinh CONFIG nhap -> admin sua trong form quan tri -> duyet."),
    ("Phe duyet/deploy: ", "co trang thai DB (draft -> pending_review -> published); trang hien ngay sau khi duyet."),
    ("Code chi can: ", "cac 'formula primitive' dung chung; them truong moi chi la them du lieu config."),
])
add_why(doc,
    "Khong chon FULL-AUTO (AI tu publish) vi do chinh xac qua quan trong - mot loi cong thuc co the "
    "khien hang nghin hoc sinh tinh sai diem. Khong chon CODEGEN (AI sinh code) vi cham va khong an "
    "toan khi chay production. Hybrid lay duoc toc do cua AI (du thao nhanh) nhung van giu con nguoi "
    "chot o khau cuoi (do chinh xac)."
)
add_why(doc,
    "Chon publish bang CO DB thay vi Git PR cho phan config vi: (1) khop pattern san co trong du an "
    "(advisor_weight_contributions.status, news_articles.published); (2) rollback chi la toggle; "
    "(3) khong can build/redeploy nen khong downtime. Git PR chi can khi gap cong thuc hoan toan moi "
    "chua co primitive - luc do dev them 1 primitive (viec hiem, lam 1 lan dung cho moi truong)."
)

# ----- 5. So sanh 3 phuong an -----
h1(doc, "5. So sanh cac phuong an (va vi sao loai)")
table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, t in enumerate(["Phuong an", "Uu diem", "Nhuoc diem", "Ket luan"]):
    hdr[i].text = t
    for pp in hdr[i].paragraphs:
        for r in pp.runs:
            r.bold = True
            r.font.size = Pt(10)
    set_cell_bg(hdr[i], "DEEAF6")

rows = [
    ("A. Full-auto config\n(AI parse -> tu publish)",
     "Nhanh nhat, gan nhu khong can nguoi.",
     "Rui ro sai cong thuc rat cao; khong ai chiu trach nhiem so lieu.",
     "LOAI - qua rui ro voi du lieu xet tuyen."),
    ("B. Hybrid (KHUYEN NGHI)\nAI nhap -> admin sua -> publish co DB",
     "Nhanh nho AI; con nguoi chot do chinh xac; khong can redeploy; rollback de.",
     "Phai xay them trang quan tri va engine config-driven.",
     "CHON - can bang tot nhat giua toc do va an toan."),
    ("C. Codegen theo PDF\nAI/dev sinh module code -> review Git",
     "Linh hoat tuyet doi, dung tinh than PDF.",
     "Cham, can dev cho moi truong, kho 'upload PDF la co'.",
     "LOAI lam mac dinh - chi dung cho cong thuc dac biet."),
]
for r in rows:
    cells = table.add_row().cells
    for i, v in enumerate(r):
        cells[i].text = v
        for pp in cells[i].paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(9.5)

# ----- 6. Kien truc -----
h1(doc, "6. Kien truc de xuat (config-driven, tang dan)")
doc.add_paragraph("Luong hoat dong tong the:")
add_code(doc,
    "Admin upload PDF\n"
    "  -> luu Supabase Storage (bucket: admission-pdfs)\n"
    "  -> POST /api/admin/admission/extract  (Gemini multimodal doc PDF -> JSON config nhap)\n"
    "  -> luu bang admission_configs (status = draft)\n"
    "  -> /admin/admission: admin review + live preview + sua\n"
    "  -> POST publish (admin only) -> status = published\n"
    "  -> Generic config-driven engine doc config\n"
    "  -> trang /unimap/[code] va /scoring?school=CODE TU DONG xuat hien"
)
doc.add_paragraph("Anh xa nguyen tac tach lop tu PDF sang kien truc moi:")
add_bullets(doc, [
    ("shared/ = ", "engine tong + cac formula primitive (weighted-combination, tsa-scale, sat-scale, cert-conversion, priority-bonus) + validate theo schema."),
    ("schools/ = ", "chi la dong config trong DB (khong con folder code moi truong cho truong moi)."),
    ("Truong cu (HUST/FTU/UET) = ", "giu nguyen module code hien tai de khong vo tinh nang dang chay; chi truong moi them-qua-PDF di duong config-driven. Migrate dan sau."),
])
add_why(doc,
    "Giu code truong cu va chi ap dung config-driven cho truong moi giup tranh mot dot refactor lon "
    "rui ro. Day cung dung loi khuyen trong PDF: 'khong gop refactor lon voi feature moi trong cung mot PR'."
)

# ----- 7. Cac buoc trien khai chi tiet -----
h1(doc, "7. Cac buoc trien khai chi tiet")

steps = [
    ("Buoc 1 - Thiet ke config-schema",
     [("Lam gi: ", "Dinh nghia cau truc config (TS type / Zod) mo ta mot truong: danh sach phuong thuc xet tuyen (methods), to hop mon (subject combinations), loai cong thuc (primitive), benchmark, diem uu tien, bang qui doi chung chi.")],
     "Schema la 'hop dong' chung giua AI, DB, engine va UI. Co schema chuan thi AI biet phai trich xuat ra cai gi, DB biet luu gi, UI biet render form gi. Day la nen tang cho moi buoc sau, lam truoc tien."),

    ("Buoc 2 - Generic engine + primitives + test",
     [("Lam gi: ", "Viet engine tong (interpreter) trong src/lib/admission-engine/generic/ cung cac primitive: weighted-combination (cong diem co he so), tsa-scale, sat-scale, cert-conversion (qui doi IELTS/chung chi), priority-bonus (diem uu tien). Kem unit test.")],
     "Day la 'bo nao' tinh diem dung chung cho moi truong config-driven. Tach primitive de tai su dung: truong moi chi can ghep cac primitive co san qua config. Bat buoc co test (input hop le / thieu field / sai kieu / output chuan / regression) vi sai so o day anh huong truc tiep ket qua hoc sinh - dung dung checklist engine cua PDF."),

    ("Buoc 3 - Noi long SchoolCode + fallback generic",
     [("Lam gi: ", "Trong src/lib/admission-engine/core/types.ts doi schoolCode tu union cung sang string. Trong registry.ts/engine.ts them fallback: khi khong tim thay module code tinh, load config published tu DB va chay generic engine. Cap nhat app/api/admission/calculate/route.ts bo whitelist cung.")],
     "Hien tai SchoolCode la union cung ('HUST'|'FTU'|...). Day la nut that lon nhat: khong the them truong moi neu khong sua type va code. Noi long thanh string + fallback generic cho phep truong moi chay ma KHONG dong vao code - dung muc tieu 'upload PDF la co'."),

    ("Buoc 4 - Database migration + Storage",
     [("Lam gi: ", "Tao migration moi (npx supabase db diff) cho bang admission_configs (school_code, year, config jsonb, status, source_pdf_url, created_by, reviewed_by, version) + bucket Storage admission-pdfs. Theo pattern status cua advisor_weight_contributions.")],
     "Config va PDF can noi luu ben vung va co trang thai duyet. Dung jsonb cho config de linh hoat. Co cot status la trai tim cua co che phe duyet. Versioning theo year vi de an doi theo tung nam - can giu lich su, khong de an moi de len an cu."),

    ("Buoc 5 - API trich xuat bang AI (extract)",
     [("Lam gi: ", "POST /api/admin/admission/extract: nhan PDF, goi Gemini (@google/genai da co san) voi structured output theo config-schema -> tra ve config nhap.")],
     "Day la phan 'tu dong' giup tiet kiem cong nhap tay. Dung Gemini multimodal doc THANG anh/PDF (vi PDF de an thuong la anh scan - chinh demo-manual.pdf cung la anh). Structured output bat AI tra dung dinh dang schema, giam loi. Nhung ket qua chi la NHAP, khong bao gio tu publish."),

    ("Buoc 6 - API admin (configs + publish)",
     [("Lam gi: ", "API admin-only: GET/POST /api/admin/admission/configs (list/luu/cap nhat draft) va POST /api/admin/admission/publish (doi status sang published). Bao ve bang isAdminRole theo mau app/api/advisor/apply-changes/route.ts.")],
     "Tach API tao/sua (draft) khoi API publish de phan quyen ro rang va co diem kiem soat duy nhat truoc khi len web chinh. Tai su dung mau requireAdmin san co de nhat quan ve bao mat."),

    ("Buoc 7 - Trang quan tri /admin/admission",
     [("Lam gi: ", "Trang upload PDF; hien PDF canh form config sua duoc; LIVE PREVIEW calculator bang config nhap; nut Duyet & Publish. Tham chieu component admin dang mo coi components/zpath/AdvisorApplyChangesClient.tsx.")],
     "Day la noi con nguoi 'chot' do chinh xac. Live preview canh PDF goc giup admin doi chieu nhanh, phat hien AI trich xuat sai truoc khi publish. Day chinh la lop an toan quan trong nhat cua ca tinh nang."),

    ("Buoc 8 - UI calculator generic + tu hien truong moi",
     [("Lam gi: ", "Them nhanh GenericConfigCalculator trong src/components/admission/AdmissionCalculatorSection.tsx (render form theo schema tu config), giu nhanh HUST/FTU/UET cu. Cho lib/unimap-visible-schools.ts doc them truong co config published de /unimap/[code] va /scoring tu xuat hien.")],
     "Buoc cuoi de trang truong moi that su xuat hien tren web chinh ngay sau khi publish. Render form theo schema (config-driven UI) la ly do mot truong moi khong can viet UI rieng. Giu nhanh truong cu de khong vo tinh nang dang chay."),
]

for i, (titletxt, what, why) in enumerate(steps, start=1):
    h2(doc, titletxt)
    for w in what:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(w[0]); r.bold = True
        p.add_run(w[1])
    add_why(doc, why)

# ----- 8. Rui ro -----
h1(doc, "8. Rui ro va luu y quan trong")
add_bullets(doc, [
    ("PDF thuong la anh scan: ", "chinh demo-manual.pdf cung la anh. Nen dung Gemini multimodal doc thang anh/PDF, khong phu thuoc text-extract thuan."),
    ("Khong bao gio auto-publish: ", "AI chi tao nhap; admin bat buoc xac nhan tung cong thuc/qui doi."),
    ("Bat buoc test generic interpreter: ", "theo checklist engine cua PDF (hop le / thieu field / sai kieu / output chuan / regression)."),
    ("Khong refactor truong cu trong cung PR: ", "giu HUST/FTU/UET nguyen ven, chi them duong config-driven cho truong moi."),
])

# ----- 9. Quyet dinh can chot -----
h1(doc, "9. Cac quyet dinh can ban chot truoc khi code")
add_bullets(doc, [
    "Dong y huong config-driven cho truong moi, giu code cu cho HUST/FTU/UET? (khuyen nghi: CO)",
    "Dong y publish bang co DB thay vi Git PR cho phan config? (khuyen nghi: CO)",
    "Co can versioning theo nam (year) ngay tu dau? (khuyen nghi: CO, vi de an doi theo nam)",
])

# ----- 10. Ket luan -----
h1(doc, "10. Ket luan")
doc.add_paragraph(
    "Phuong an toi uu la di theo huong config-driven co AI ho tro va con nguoi phe duyet. "
    "Cach nay giu duoc tinh than tach lop tot cua file demo-manual.pdf, dong thoi giai duoc bai toan "
    "that su cua ban: bien mot file PDF de an thanh mot trang tinh diem chay duoc, an toan, va chi len "
    "web chinh khi da duoc duyet - tat ca ma khong can deploy lai moi khi them truong."
)

out = r"e:\zpath\zpath\docs\KeHoach-AutoAdmission-TuPDF.docx"
doc.save(out)
print("SAVED:", out)
