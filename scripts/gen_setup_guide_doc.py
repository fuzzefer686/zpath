# -*- coding: utf-8 -*-
"""Generate an English Word doc: how to run & use the auto-admission feature."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x55, 0x55, 0x55)
CODE_BG = "F2F2F2"


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)
    return p


def add_note(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10.5)
    b = p.add_run(text)
    b.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# ----- Cover -----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("SETUP & USAGE GUIDE\nAuto-generate per-school admission calculators from PDF")
trun.bold = True
trun.font.size = Pt(19)
trun.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = sub.add_run("ZPATH - AI Career Guidance Platform")
srun.italic = True
srun.font.size = Pt(12)
srun.font.color.rgb = MUTED

doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    "The feature is already implemented in the codebase. The steps below are what "
    "YOU need to do to run the database, configure the environment, and test the "
    "full flow: upload an admission-scheme PDF, let AI draft a scoring config, "
    "review it, approve it, and see the school's calculator appear on the site."
).font.size = Pt(11)

# ----- Prerequisites -----
h1(doc, "Prerequisites")
add_bullets(doc, [
    "Docker Desktop installed and running.",
    "Node.js 22.x and npm installed.",
    "A Google Gemini API key (Google AI Studio) OR a Vertex service account.",
    "Project located at e:\\zpath\\zpath.",
])

# ----- Step 1 -----
h1(doc, "Step 1 - Start the local database")
doc.add_paragraph("Open Docker Desktop first, then from the project folder:")
add_code(doc, "cd e:\\zpath\\zpath\nnpx supabase start")
add_note(doc, "Note",
    "This command prints the API URL, the service_role key and the anon key. "
    "You will need these values in Step 3.")

# ----- Step 2 -----
h1(doc, "Step 2 - Apply the migration (table + storage bucket)")
doc.add_paragraph(
    "For local dev, run all pending migrations. For production/preview on "
    "Vercel, the project owner must run 'npx supabase db push' on the linked "
    "Supabase project after merge. The migration that provisions generate-admission "
    "infrastructure is 20260613140000_ensure_admission_generate_system.sql "
    "(idempotent; safe even if an older admission_configs migration was never applied).")
add_code(doc, "npx supabase migration up")
doc.add_paragraph(
    "If the migration state is out of sync locally, run the command below instead "
    "(WARNING: this wipes all local data and re-runs every migration + seed):")
add_code(doc, "npx supabase db reset")
doc.add_paragraph(
    "Then verify in Supabase Studio (the URL printed by 'supabase start'): you "
    "should see the table 'admission_configs' and the storage bucket "
    "'admission-pdfs'.")
doc.add_paragraph(
    "On Vercel preview/production, keep PDF uploads under ~4 MB (platform "
    "request body limit). Compress large admission PDFs before uploading.")

# ----- Step 3 -----
h1(doc, "Step 3 - Configure .env.local")
doc.add_paragraph(
    "Open e:\\zpath\\zpath\\.env.local and make sure these variables exist "
    "(use the values printed in Step 1):")
add_code(doc,
    "SUPABASE_URL=http://127.0.0.1:54321\n"
    "NEXT_PUBLIC_SUPABASE_URL=http://127.0.0.1:54321\n"
    "NEXT_PUBLIC_SUPABASE_ANON_KEY=<anon key>\n"
    "SUPABASE_SERVICE_ROLE_KEY=<service_role key>\n"
    "\n"
    "# AI used to read the PDF (pick ONE option)\n"
    "GEMINI_API_KEY=<key from Google AI Studio>\n"
    "# or Vertex: GOOGLE_APPLICATION_CREDENTIALS_JSON / _BASE64\n"
    "\n"
    "# Make your account an admin\n"
    "ADMIN_EMAILS=your-email@gmail.com")
add_note(doc, "Required",
    "Two things are mandatory for this feature: SUPABASE_SERVICE_ROLE_KEY "
    "(reads/writes configs + storage) and a Gemini key (reads the PDF). Without "
    "the Gemini key the 'Extract with AI' button fails, but you can still paste "
    "the JSON config manually.")

# ----- Step 4 -----
h1(doc, "Step 4 - Run the app and sign in as admin")
add_code(doc, "npm run dev")
doc.add_paragraph(
    "Open http://localhost:3001. Register/sign in with the exact email you put in "
    "ADMIN_EMAILS (or the default admin email). The navbar should show the "
    "\"Quan tri vien\" (Administrator) badge.")

# ----- Step 5 -----
h1(doc, "Step 5 - Use the admin page")
doc.add_paragraph(
    "Go to http://localhost:3001/admin/admission. You have a real admission-scheme "
    "PDF to test with in e:\\zpath\\ftu\\ (Foreign Trade University). The flow:")
add_bullets(doc, [
    ("Section 1 - Upload & extract: ", "enter School code (e.g. VNU), School name, Year 2026, choose the PDF file, then click \"Trich xuat bang AI\" (Extract with AI)."),
    ("Section 2 - Config (JSON): ", "the AI fills the JSON box. Read it carefully and correct any formulas/conversions. A green \"Cau hinh hop le\" (Config valid) line means it matches the schema."),
    ("Section 3 - Preview: ", "the calculator renders live on the right; enter sample scores to verify the result."),
    ("Approve: ", "click \"Luu ban nhap\" (Save draft), then \"Duyet & Publish\" (Approve & Publish)."),
])

# ----- Step 6 -----
h1(doc, "Step 6 - Verify on the main site")
doc.add_paragraph(
    "Open http://localhost:3001/scoring?school=VNU (replace VNU with the school "
    "code you just published). The new school must appear in the school selector "
    "and compute scores - with NO redeploy required.")

# ----- Step 7 -----
h1(doc, "Step 7 - Run the engine tests (recommended)")
add_code(doc, "npx tsx tests/genericAdmissionEngine.test.ts")
doc.add_paragraph("It must print: all assertions passed.")

# ----- Step 8 -----
h1(doc, "Step 8 - Commit using Gitflow")
doc.add_paragraph(
    "Per README.md, never commit directly to main. Create a feature branch off "
    "develop:")
add_code(doc,
    "git checkout develop\n"
    "git pull origin develop\n"
    "git checkout -b feature/auto-admission-from-pdf")
doc.add_paragraph(
    "Then open a Pull Request into develop. (The assistant does not auto-commit; "
    "ask it explicitly when you want the changes committed.)")

# ----- Troubleshooting -----
h1(doc, "Troubleshooting")
add_bullets(doc, [
    ("403 on /admin/admission: ", "your account is not an admin. Check ADMIN_EMAILS and sign in with that email."),
    ("\"Extract with AI\" fails: ", "Gemini key missing/invalid. Set GEMINI_API_KEY, or paste the config JSON manually."),
    ("Bucket not found on Vercel: ", "owner must run supabase db push so 20260613140000_ensure_admission_generate_system.sql is applied on the cloud project."),
    ("PDF too large on Vercel: ", "keep files under ~4 MB or compress the PDF."),
    ("Missing SUPABASE_SERVICE_ROLE_KEY error: ", "set it in .env.local and restart npm run dev."),
    ("School not showing on /scoring: ", "make sure the config status is 'published', not just 'draft'."),
    ("PDF is a scanned image: ", "that is fine - Gemini reads the PDF/images directly; no separate OCR needed."),
])

out = r"e:\zpath\zpath\docs\Setup-Guide-AutoAdmission-EN.docx"
doc.save(out)
print("SAVED:", out)
